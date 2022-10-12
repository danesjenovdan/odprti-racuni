from nvo.models import RevenueCategory, ExpensesCategory
from docx import Document, oxml, opc
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from datetime import datetime
import markdown
from htmldocx import HtmlToDocx
from babel.numbers import format_decimal

from django.utils.text import slugify
from django.http.response import HttpResponse


revenue_data = {
    'model': RevenueCategory,
    'nodes': [
        {'name': 'Skupni prihodki', 'order': 1, 'allow_additional_name': False},
        {'name': 'Tržna dejavnost', 'parent': 0, 'order': 2, 'allow_additional_name': False},
        {'name': 'Sponzorstva', 'order': 3, 'parent': 1, 'allow_additional_name': False},
        {'name': 'Produkti in storitve', 'order': 4, 'parent': 1, 'allow_additional_name': False},
        {'name': 'Donacije in subvencije', 'parent': 0, 'order': 5, 'allow_additional_name': False},
        {'name': 'Donacije fizičnih oseb', 'order': 6, 'parent': 4, 'allow_additional_name': False},
        {'name': 'Javna sredstva', 'order': 7, 'parent': 4, 'allow_additional_name': False},
        {'name': 'Državna sredstva', 'order': 8, 'parent': 6, 'allow_additional_name': False},
        {'name': 'Občinska sredstva', 'order': 9, 'parent': 6, 'allow_additional_name': False},
        {'name': 'Evropska sredstva', 'order': 10, 'parent': 6, 'allow_additional_name': False},
        {'name': 'Sredstva drugih držav', 'order': 11, 'parent': 6, 'allow_additional_name': False},
        {'name': 'Zasebne fundacije', 'order': 12, 'parent': 4, 'allow_additional_name': False},
        {'name': 'Članarine', 'order': 13, 'parent': 4, 'allow_additional_name': False},
        {'name': 'Drugo', 'order': 14, 'instructions': 'Vnesi dodatno ime ki se bo prikazovalo poleg texta Drugo <vaš text>', 'parent': 4, 'allow_additional_name': True},
        {'name': 'Drugi prihodki', 'order': 15, 'parent': 0, 'allow_additional_name': True, 'instructions': 'Vnesi dodatno ime ki se bo prikazovalo'},
]}

expenses_data = {
    'model': ExpensesCategory,
    'nodes': [
        {'name': 'Skupni odhodki', 'order': 1, 'allow_additional_name': False},
        {'name': 'Materiali in storitve', 'parent': 0,'order': 2, 'allow_additional_name': False},
        {'name': 'Redno delovanje', 'order': 3, 'parent': 1, 'allow_additional_name': False},
        {'name': 'Zunanje storitve in izvajalci', 'order': 4, 'parent': 1, 'allow_additional_name': False},
        {'name': 'Potni stroški', 'order': 5, 'parent': 1, 'allow_additional_name': False},
        {'name': 'Drugo', 'order': 6, 'parent': 1, 'allow_additional_name': True},
        {'name': 'Delo', 'order': 7, 'parent': 0, 'allow_additional_name': False},
        {'name': 'Plače', 'order': 8, 'parent': 6, 'allow_additional_name': False},
        {'name': 'Pokojninska in druga socialna zavarovanja', 'order': 9, 'parent': 6, 'allow_additional_name': False},
        {'name': 'Drugo', 'order': 10, 'parent': 6, 'allow_additional_name': False},
        {'name': 'Drugi odhodki', 'order': 11, 'parent': 0, 'allow_additional_name': True, 'instructions': 'Vnesi dodatno ime ki se bo prikazovalo'},
]}
financial_data = [revenue_data, expenses_data]

def create_financial_tree(year_id, organization_id):
    for item in financial_data:
        create_financial_tree_for_model(
            item['nodes'],
            item['model'],
            year_id,
            organization_id
        )

def create_financial_tree_for_model(data, model, year_id, organization_id):
    for element in data:
        if 'parent' in element.keys():
            parent = parent=data[element['parent']]['object']
        else:
            parent = None
        obj = model(
            name=element['name'],
            organization_id=organization_id,
            year_id=year_id,
            order=element['order'],
            parent=parent,
            allow_additional_name=element['allow_additional_name'])
        obj.save()
        element['object'] = obj
        if 'instructions' in element.keys():
            obj.instructions=element['instructions']
        obj.save()

def clean_chart_data(data):
    '''
    Recursively cleans financial data to be able to draw the chart.
    '''
    if len(data['children']) > 0:
        data['value'] = None
        data['children'] = [clean_chart_data(child) for child in data['children']]
    elif 'amount' in data.keys():
        data['value'] = data['amount']

    return data


class ExportNVOData(object):
    def __init__(self, organization, year):
        self.organization = organization
        self.year = year
        self.document = Document()
        self.set_arial()

    def get_response(self):
        header = self.document.sections[0].header

        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        logo_run = paragraph.add_run()
        logo_run.add_picture("static_files/images/or.png", width=Inches(1.5))
        logo_run.bottom_margin = Inches(1)

        self.insertHR(paragraph)

        self.add_heading('Osebna izkaznica\n')

        self.write_paragraph(
            title=self.organization.name,
            lines=[
                self.organization.address,
                self.organization.post_number
            ]
        )

        self.write_paragraph(
            lines=[
                f'Davčna številka: {self.organization.post_number}',
                f'Matična številka: {self.organization.post_number}'
            ]
        )

        self.write_paragraph(
            lines=[
                f'Kontakt: {self.organization.email}, {self.organization.phone_number}',
                f'Zastopnik_ca: {self.organization.representative}'
            ]
        )

        self.write_paragraph(
            lines=[
                f'TRR: {self.organization.trr}'
            ]
        )

        if self.organization.is_for_the_public_good:
            self.write_paragraph(
                lines=[
                    f'Organizacija ima status humanitarne organizacije in organizacije v javnem interesu na področju {self.organization.is_for_the_public_good}'
                ],
                add_newline=True
            )

        p = self.document.add_paragraph()
        p.add_run(f'Letna poročila\n').bold = True
        for doc in self.organization.documents.all(): # TODO filter by year
            hyperlink = self.add_hyperlink(
                p,
                f'{doc.file.url}',
                f'{doc.category.name}\n', # text
            )

        people_stats = self.organization.people.first().get_statistics() # TODO filter by name
        if people_stats:
            self.write_paragraph(
                title='Ljudje',
                lines=[
                    f'Število redno zaposlenih: ',
                    f'Redno zaposleni po spolu: {format_decimal(people_stats["men"], locale="sl_SI")} % moških, {format_decimal(people_stats["women"], locale="sl_SI")} % žensk, {format_decimal(people_stats["nonbinary"], locale="sl_SI")} % nebinarnih oseb'
                ],
                add_newline=True
            )

        payment_ratios = self.organization.payment_ratios.first().get_statistics()
        if payment_ratios:
            self.write_paragraph(
                title='Plačna razmerja',
                lines=[
                    f'Najvišja plača: Najnižja plača: {format_decimal(payment_ratios["highest_absolute"], locale="sl_SI")}:{format_decimal(payment_ratios["lowest"], locale="sl_SI")}',
                    f'Najvišja plača: Povprečna plača: {format_decimal(payment_ratios["highest"], locale="sl_SI")}:{format_decimal(payment_ratios["average"], locale="sl_SI")}'
                ]
            )
        self.document.add_page_break()

        # FINANCE
        self.add_heading('Finance\n')

        reveues = self.organization.revenuecategory_related.filter(year=self.year).exclude(level=0)
        total_income = RevenueCategory.objects.get(year=self.year, organization=self.organization, level=0)
        if total_income.amount:
            self.write_paragraph(
                title=f'Skupni prihodki: {format_decimal(total_income.amount, locale="sl_SI")} EUR\n',
                lines=[
                    f'{"    " * revenue.level} {revenue.name}: {format_decimal(revenue.amount, locale="sl_SI")} EUR' for revenue in reveues if revenue.amount > 0
                ]
            )

        expenses = self.organization.expensescategory_related.filter(year=self.year).exclude(level=0)
        total_expense = ExpensesCategory.objects.get(year=self.year, organization=self.organization, level=0)
        if total_expense.amount:
            self.write_paragraph(
                title=f'\nSkupni odhodki: {format_decimal(total_expense.amount, locale="sl_SI")} EUR\n',
                lines=[
                    f'{"    " * expense.level} {expense.name}: {format_decimal(expense.amount, locale="sl_SI")} EUR' for expense in expenses if expense.amount > 0
                ]
            )

        finance = self.organization.finances.get(year=self.year)

        if finance.payments_project_partners:
            self.write_paragraph(
                lines=[
                    f'Izplačila projektnim partnerjem: {format_decimal(finance.payments_project_partners, locale="sl_SI")} EUR'
                ],
                hr=True
            )

        if finance.payment_state_budget:
            self.write_paragraph(
                lines=[
                    f'Vplačila v proračun RS: {format_decimal(finance.payment_state_budget, locale="sl_SI")} EUR',
                    f'Razlika med vplačanimi in pridobljenimi proračunskimi sredstvi v RS: {format_decimal(finance.difference_state_budget, locale="sl_SI")} EUR'
                ],
                hr=True
            )

        self.document.add_page_break()

        # PROJECTS
        self.add_heading('Projekti\n')

        projects = self.year.get_projects().filter(organization=self.organization)
        for project in projects:
            self.write_projects_paragraph(project)
            self.document.add_page_break()



        # DONATIONS
        self.add_heading('Donacije\n')

        donations = self.organization.donations.get(year=self.year)
        if donations.personal_donations_amount:
            self.write_paragraph(
                title='Donacije fizičnih oseb',
                lines=[
                    f'Višina zbranih donacij fizičnih oseb: {format_decimal(donations.personal_donations_amount, locale="sl_SI")} EUR',
                    f'Število donatorjev: {donations.number_of_personal_donations}'
                ]
            )
            self.write_list_of_donors(donations.personal_donators.all())

        donations = self.organization.donations.get(year=self.year)
        if donations.organization_donations_amount:
            self.write_paragraph(
                title='Donacije pravnih oseb',
                lines=[
                    f'Višina zbranih donacij fizičnih oseb: {format_decimal(donations.organization_donations_amount, locale="sl_SI")} EUR',
                    f'Število donatorjev: {donations.number_of_organization_donations}'
                ],
                hr=True
            )
            self.write_list_of_donors(donations.organiaztion_donators.all())

        if donations.one_percent_income_tax:
            self.insertDashes()
            p = self.document.add_paragraph()
            p.add_run('Višina zbranih donacij z 1% dohodnine: ').bold = True
            p.add_run(f'{format_decimal(donations.one_percent_income_tax, locale="sl_SI")} EUR')


        if donations.purpose_of_donations:
            self.write_paragraph(
                title='Kako smo porabili zbrane donacije',
                lines=[
                    donations.purpose_of_donations,

                ],
                hr=True
            )

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={slugify(self.organization.name)}_{self.year.name}.docx'
        self.document.save(response)

        return response

    def add_heading(self, text, level=1):
        heading = self.document.add_heading('', level=level).add_run(text)
        heading.font.color.rgb = RGBColor(0, 0, 0)

    def write_paragraph(self, title=None, lines=[], hr=False, add_newline=False):
        if hr:
            self.insertDashes()
        p = self.document.add_paragraph()
        if title:
            p.add_run(f'{title}\n').bold = True
        for i, line in enumerate(lines):
            if i < len(lines)-1 or add_newline:
                text = f'{line}\n'
            else:
                text = f'{line}'
            p.add_run(text)

    def write_list_of_donors(self, donors):
        p = self.document.add_paragraph()
        p.add_run('Poimenski seznam donatorjev')
        p = self.document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        for donor in donors:
            p.add_run(f'{donor.name}: {format_decimal(donor.amount, locale="sl_SI")} EUR\n')

    def write_projects_paragraph(self, project):
        duration = project.duration
        months = duration["months"]
        days = duration["days"]
        if days:
            days = f' in {days} dni'
        else:
            days = ''
        self.add_heading(project.name, level=2)
        p = self.document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(f'Vrednost projekta (delež organizacije / celota): {format_decimal(project.organization_share, locale="sl_SI")} EUR / {format_decimal(project.value, locale="sl_SI")} EUR\n')
        p.add_run(f'Trajanje projekta: {months} mesecev{days}, {self.get_formated_date(project.start_date)}–{self.get_formated_date(project.end_date)}\n')
        financers = project.financers.all()
        if financers:
            p.add_run(f'Financerji: {", ".join([financer.name for financer in financers])}\n')
        cofinancers = project.cofinancers.all()
        if cofinancers:
            p.add_run(f'Sofinancerji: {", ".join([cofinancer.name for cofinancer in cofinancers])}\n')
        partners = project.partners.all()
        if partners:
            p.add_run(f'Partnerji: {", ".join([partner.name for partner in partners])}\n')
        donators = project.donators.all()
        if donators:
            p.add_run(f'Donatorji: {", ".join([donator.name for donator in donators])}\n')

        if project.self_money:
            p.add_run(f'Lastni vložek: {format_decimal(project.self_money, locale="sl_SI")} EUR\n')
        if project.link:
            self.add_hyperlink(p, project.link, 'Projektno spletno mesto')

        self.added_indented_markdown(project.description)

        p = self.document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(f'Predvideni rezultati in učinki:')
        self.added_indented_markdown(project.outcomes_and_impacts, 1)

        if project.icons:
            self.insertDashes(indent=0.5)
            p = self.document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            r = p.add_run()
            for icon in project.icons:
                r.add_picture(icon.path, height=Inches(1))



    def added_indented_markdown(self, markdown_text, indention=0.5):
        new_parser = HtmlToDocx()
        num_paragraphs = len(self.document.paragraphs)
        new_parser.add_html_to_document(markdown.markdown(markdown_text), self.document)

        for i in range(num_paragraphs, len(self.document.paragraphs)):
            self.document.paragraphs[i].paragraph_format.left_indent = Inches(indention)

    def get_formated_date(self, date):
        return datetime.strftime(date, '%d. %m. %Y')

    def add_hyperlink(self, paragraph, url, text, color='0000EE', underline=False):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = oxml.shared.OxmlElement('w:rPr')

        # Add color if it is given
        if not color is None:
            c = oxml.shared.OxmlElement('w:color')
            c.set(oxml.shared.qn('w:val'), color)
            rPr.append(c)

        # Remove underlining if it is requested
        if not underline:
            u = oxml.shared.OxmlElement('w:u')
            u.set(oxml.shared.qn('w:val'), 'none')
            rPr.append(u)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    def set_arial(self):
        font = self.document.styles['Normal'].font
        font.name = 'Arial'

    def insertHR(self, paragraph):
        p = paragraph._p  # p is the <w:p> XML element
        pPr = p.get_or_add_pPr()
        pBdr = oxml.shared.OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = oxml.shared.OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '2')
        bottom.set(qn('w:space'), '10')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def insertDashes(self, indent=None):
        p = self.document.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(0.5)
        p.add_run('---------------------')
